from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parent
ARCHIVE_DIR = ROOT_DIR.parent.parent / "归档"
OUTPUT_PATH = ROOT_DIR / "VibecodingKeyboard_macOS使用说明书_流程版_V1.0.0.docx"


def set_east_asia_font(run_or_style, font_name: str) -> None:
    r_pr = run_or_style.element.rPr
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    set_east_asia_font(normal, "微软雅黑")

    for style_name, size, bold in [
        ("Title", 24, True),
        ("Subtitle", 13, False),
        ("Heading 1", 16, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11.5, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        set_east_asia_font(style, "微软雅黑")

    if "Tip" not in doc.styles:
        tip_style = doc.styles.add_style("Tip", WD_STYLE_TYPE.PARAGRAPH)
        tip_style.base_style = doc.styles["Normal"]
        tip_style.font.name = "Calibri"
        tip_style.font.size = Pt(10.5)
        tip_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        set_east_asia_font(tip_style, "微软雅黑")


def add_text(doc: Document, text: str, style: str = "Normal", *, bold: bool = False,
             color: RGBColor | None = None, align: WD_ALIGN_PARAGRAPH | None = None,
             space_after: float = 6) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "Calibri"
    set_east_asia_font(run, "微软雅黑")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Calibri"
        set_east_asia_font(run, "微软雅黑")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Calibri"
        set_east_asia_font(run, "微软雅黑")


def add_tip(doc: Document, text: str) -> None:
    add_text(doc, f"提示：{text}", style="Tip", color=RGBColor(0x1F, 0x4E, 0x79), space_after=8)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_screenshot(doc: Document, image_name: str, caption: str, note: str | None = None) -> None:
    image_path = ARCHIVE_DIR / image_name
    if not image_path.exists():
        add_tip(doc, f"截图缺失：{image_name}")
        return

    doc.add_picture(str(image_path), width=Inches(6.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run(caption)
    run.bold = True
    run.font.name = "Calibri"
    set_east_asia_font(run, "微软雅黑")
    if note:
        add_text(doc, note)


def add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(" ")
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9D9D9")
    borders.append(bottom)
    p_pr.append(borders)


def build_manual(doc: Document) -> None:
    add_text(
        doc,
        "Vibecoding Keyboard macOS 使用说明书",
        style="Title",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_text(
        doc,
        "流程版 V1.0.0 | 参考 Vibecoding 使用文档 V1.2.1 编排 | 基于归档截图整理",
        style="Subtitle",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x66, 0x66, 0x66),
        space_after=4,
    )
    add_text(
        doc,
        "更新时间：2026-04-13",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x66, 0x66, 0x66),
        space_after=12,
    )
    add_text(
        doc,
        "本文档单独面向 macOS 版本使用场景编写，目标是让用户只看这一份文档，就能按截图顺序完成安装、首次启动、Hook 安装、系统授权、语音输入、设备连接与模式配置。",
        space_after=8,
    )
    add_text(
        doc,
        "即使某些概念在 Windows 说明书里已经出现过，这里也会再次完整说明，避免用户在 macOS 流程里来回对照两份文档。",
        space_after=10,
    )
    add_divider(doc)

    add_text(doc, "免责说明", style="Heading 1", space_after=8)
    disclaimer = [
        "本文档提供有关南京市锦心湾科技有限责任公司相关产品、方案及技术资料的信息。除非另有书面约定，本文档内容不构成任何形式的明示或暗示保证，也不视为授予任何知识产权许可或其他权利。",
        "本文档所载的文字、图片、结构设计、技术方案、软件程序及其他相关内容，其著作权、专利权、商标权、商业秘密及其他合法权益均归南京市锦心湾科技有限责任公司所有，并受中华人民共和国相关法律法规保护。未经南京市锦心湾科技有限责任公司事先书面授权，任何单位或个人不得以任何方式对本文档全部或部分内容进行复制、转载、引用、修改、传播、披露、反向工程、出售或用于其他商业用途。",
        "本文档所涉及的产品功能、技术参数、接口定义、应用场景及实施方案等内容，可能因产品迭代、版本升级或业务调整而发生变更。南京市锦心湾科技有限责任公司保留在不另行通知的情况下对本文档内容进行补充、修订、更新或撤回的权利。",
        "在订购、部署或使用相关产品与服务前，请与南京市锦心湾科技有限责任公司联系，以获取最新版本的技术说明、产品资料及商务信息。",
        "南京市锦心湾科技有限责任公司保留所有权利。",
    ]
    for para in disclaimer:
        add_text(doc, para)

    add_text(doc, "一、文档适用范围与阅读顺序", style="Heading 1", space_after=8)
    add_text(
        doc,
        "本文档适合以下对象直接使用：第一次在 Mac 上安装 Vibecoding Keyboard 的用户、需要按图验收流程的测试同事、需要给客户演示安装与授权流程的售后同事。",
    )
    add_text(doc, "建议按下面顺序阅读和操作：")
    add_numbered(
        doc,
        [
            "先完成键盘硬件基础操作，包括开机、关机、模式切换和蓝牙配对。",
            "把 Vibecoding Keyboard 安装到“应用程序”目录，并从应用程序中打开。",
            "按首次启动提示处理文稿访问、蓝牙权限和 Hook 管理器。",
            "点击“启动语音输入”，继续完成麦克风、输入监控和辅助功能权限。",
            "等语音状态变为“语音已就绪”后，再实际测试语音输入。",
            "连接设备后进入“模式配置”页，根据需要设置快捷键、宏定义和动画。",
        ],
    )
    add_tip(
        doc,
        "Mac 上“启动语音输入”和“连接设备”是两条不同的链路。前者用于启动本地语音服务，后者用于写入键位和动画配置。语音能用，不代表设备配置链路已经连好。",
    )

    add_text(doc, "二、开箱与硬件基础操作", style="Heading 1", space_after=8)
    add_text(doc, "在正式进入软件前，先确认以下硬件操作含义：")
    add_bullets(
        doc,
        [
            "机器开机：短按电源按键，屏幕正常亮起即表示开机成功；若屏幕无反应，先给设备充电。",
            "机器关机：长按电源按键，直到指示灯变为红色后快速熄灭，即表示关机。",
            "模式切换：单击电源键即可循环切换 Mode0、Mode1、Mode2。",
            "Mode0 是默认 Vibecoding 功能模式。当前默认预设为：Key1 为 F18 语音键，Key2 为 YES，Key3 为 NO，Key4 为 Enter。",
            "Key1 对应的是键盘最左边的语音键。实际进行语音输入时，按住这个语音键说话，松开后进入识别。",
            "蓝牙配对：短按设备右下角白色按钮，让设备进入配对状态，再到 Mac 的蓝牙设置里找到“vibe code”完成连接。",
        ],
    )
    add_tip(
        doc,
        "如果你只是先体验语音输入，可以先完成软件安装与系统授权；如果你要把快捷键、宏或动画写入设备，则后面还需要在主界面顶部点击“连接”。",
    )

    add_text(doc, "三、按截图完成 Mac 安装与首次启动", style="Heading 1", space_after=8)
    add_text(
        doc,
        "下面这一节严格按归档截图的顺序串起整个 Mac 流程。你可以一边看文档，一边对照自己电脑上的实际界面逐步操作。",
        space_after=8,
    )

    add_text(doc, "3.1 安装应用到“应用程序”", style="Heading 2", space_after=8)
    add_screenshot(
        doc,
        "step01.png",
        "图 1 打开 DMG 后的安装界面",
        "双击 macOS 安装包后，会看到 Vibecoding Keyboard、Applications 和安装说明。Mac 端标准做法不是直接在 DMG 里长期运行，而是先把应用拖到“应用程序”目录。",
    )
    add_screenshot(
        doc,
        "step02.png",
        "图 2 把应用拖入 Applications",
        "把 Vibecoding Keyboard 拖到 Applications 后，再从“应用程序”里启动。这样后面做系统权限授权时，路径会稳定保持为 /Applications/Vibecoding Keyboard.app，也更符合 macOS 的权限管理逻辑。",
    )
    add_screenshot(
        doc,
        "step03.png",
        "图 3 从应用程序首次打开应用",
        "第一次打开时，系统可能先进行安全校验或显示“正在验证”之类提示，属于正常现象，等待校验完成即可。",
    )

    add_text(doc, "3.2 首次进入主界面", style="Heading 2", space_after=8)
    add_screenshot(
        doc,
        "step04.png",
        "图 4 首次进入主界面时的提示",
        "此时通常会同时看到主界面、文稿访问权限提示，以及 Hook 缺失提醒。若系统询问是否允许访问“文稿”目录，建议允许；若提示 Claude、Cursor Hook 未安装，建议按提示进入 Hook 管理器完成安装。",
    )
    add_screenshot(
        doc,
        "step05.png",
        "图 5 首次启动后的主界面状态",
        "首次打开时看到多个提示框并不代表出错，而是说明系统正在逐项完成权限和环境准备。Mac 端建议不要跳过这些步骤，后续语音输入和终端联动都会依赖它们。",
    )
    add_tip(
        doc,
        "如果当前版本在主窗口展示后会先出现“欢迎使用 Vibecoding Keyboard”的首次引导，请先阅读引导并点击“我知道了”，然后继续处理 Hook 和系统权限。这类引导只会自动出现一次。",
    )

    add_text(doc, "3.3 蓝牙权限与 Hook 安装", style="Heading 2", space_after=8)
    add_screenshot(
        doc,
        "step06.png",
        "图 6 蓝牙权限请求",
        "系统弹出蓝牙权限请求时，请选择允许。即使当前阶段你还没马上进入设备写配置，这个权限也建议在首次安装时一次性放行，避免后面连接设备时再次中断流程。",
    )
    add_screenshot(
        doc,
        "step07.png",
        "图 7 Hook 管理器中 Claude 和 Cursor 均未安装",
        "打开“Claude / Cursor Hook 管理工具”后，如果两个状态都显示“未安装”，建议分别点击安装。Hook 的作用是让终端或编辑器在配合 Vibecoding 使用时，能正确接收和联动相关能力。",
    )
    add_screenshot(
        doc,
        "step08.png",
        "图 8 Hook 安装完成",
        "当 Claude Hook 和 Cursor Hook 都显示“已安装”后，就说明这一环已经完成。截图里同时展示了写入路径，例如 Claude 的配置文件会写到 ~/.claude/settings.json，Cursor 的 Hook 会写到 ~/.cursor/hooks.json。",
    )
    add_tip(
        doc,
        "如果你当前只打算先测试语音输入，Hook 不是阻断项；但如果你要完整体验 Claude / Cursor 联动，建议第一次就把 Hook 安装完。",
    )

    add_text(doc, "3.4 进入主界面并启动语音输入", style="Heading 2", space_after=8)
    add_screenshot(
        doc,
        "step09.png",
        "图 9 回到主界面后的默认状态",
        "回到主界面后，可以先确认顶部按钮区、设备区和“模式配置”页是否正常显示。截图里能看到 Mode0 默认预设已经直接显示为 Key1 F18、Key2 YES、Key3 NO、Key4 Enter，这说明即使没有真实自定义值，界面也会给出默认显示。",
    )
    add_screenshot(
        doc,
        "step10.png",
        "图 10 点击“启动语音输入”并处理麦克风权限",
        "点击“启动语音输入”后，主界面右侧状态会进入“语音启动中”，同时系统会请求麦克风权限。这里必须允许，否则无法录音。",
    )

    add_text(doc, "3.5 完成输入监控与辅助功能授权", style="Heading 2", space_after=8)
    add_screenshot(
        doc,
        "step11.png",
        "图 11 继续完成系统授权",
        "麦克风通过后，应用通常会继续提示还缺少“输入监控”和“辅助功能”权限。这两个权限对语音输入结果回写、按键监听和系统交互都很关键。按弹窗中的按钮逐项打开对应系统设置即可。",
    )
    add_screenshot(
        doc,
        "step12.png",
        "图 12 打开“输入监控”设置页",
        "进入“系统设置 -> 隐私与安全性 -> 输入监控”后，如果列表中还没有 Vibecoding Keyboard，就需要手动添加。",
    )
    add_screenshot(
        doc,
        "step13.png",
        "图 13 先解锁权限页",
        "macOS 的隐私设置通常需要先点左下角锁并输入密码，或者用 Touch ID 解锁。没有先解锁，后续就无法添加应用。",
    )
    add_screenshot(
        doc,
        "step14.png",
        "图 14 从应用程序中添加 Vibecoding Keyboard.app",
        "在文件选择器中进入“应用程序”，找到 Vibecoding Keyboard.app 并点击“打开”。如果系统列表里已经存在此应用，也请确认开关处于开启状态。",
    )
    add_screenshot(
        doc,
        "step15.png",
        "图 15 回到应用继续下一步",
        "输入监控添加完成后，回到应用中的授权对话框，继续处理剩余权限。",
    )
    add_screenshot(
        doc,
        "step16.png",
        "图 16 打开“辅助功能”设置页",
        "接下来进入“系统设置 -> 隐私与安全性 -> 辅助功能”，操作方法和输入监控基本一致，同样需要把 Vibecoding Keyboard.app 加进去并勾选允许。",
    )
    add_screenshot(
        doc,
        "step17.png",
        "图 17 把应用添加到辅助功能",
        "如果辅助功能列表里没有 Vibecoding Keyboard，就从“应用程序”目录重新添加一次。路径同样建议保持为 /Applications/Vibecoding Keyboard.app。",
    )
    add_tip(
        doc,
        "如果系统设置列表里已经有 Vibecoding Keyboard，但仍然不生效，可以先关闭勾选再重新勾选，然后完全退出应用后再重新打开。",
    )

    add_text(doc, "3.6 返回应用等待语音就绪", style="Heading 2", space_after=8)
    add_screenshot(
        doc,
        "step18.png",
        "图 18 返回应用后等待模型加载",
        "系统权限全部处理完以后，回到主界面，语音状态通常会先显示“模型加载中”或类似阶段性提示。此时按钮已经切换成“停止语音输入”，说明语音服务已经在启动过程中。",
    )
    add_screenshot(
        doc,
        "step19.png",
        "图 19 主界面稳定后的状态",
        "当顶部状态最终进入“语音已就绪”后，就可以把光标放到任意输入框，按住键盘最左边的语音键开始说话，松开后等待识别结果输入。截图同时展示了设备已连接状态，说明此时也可以进入后续的模式配置。",
    )

    add_text(doc, "四、Mac 端日常使用流程", style="Heading 1", space_after=8)
    add_text(doc, "完成首次安装后，日常使用建议按下面两条路径理解：")
    add_bullets(
        doc,
        [
            "只做语音输入：打开应用 -> 点击“启动语音输入” -> 等状态变成“语音已就绪” -> 按住语音键说话 -> 松开后等待识别。",
            "修改设备按键或动画：打开应用 -> 确认蓝牙与设备通信链路正常 -> 点击顶部“连接” -> 进入“模式配置”页 -> 修改按键映射或动画 -> 应用到设备或保存到设备。",
        ],
    )
    add_tip(
        doc,
        "如果你点击“应用按键到设备”时提示“请先连接设备”，一般不是蓝牙坏了，而是说明上位机和设备的配置通信链路还没连上。先确认设备在线，再点顶部“连接”。",
    )

    add_text(doc, "4.1 如何判断语音是否真的可用", style="Heading 2", space_after=8)
    add_bullets(
        doc,
        [
            "顶部按钮从“启动语音输入”切换成“停止语音输入”。",
            "语音状态灯和文字最终变成绿色的“语音已就绪”。",
            "把光标放到文本输入区后，按住语音键说话，松开后会进入识别并写回文字。",
        ],
    )
    add_text(
        doc,
        "若仍停留在“语音启动中”“模型加载中”或权限类提示，请优先回头检查麦克风、输入监控和辅助功能三项系统授权是否都已开启。",
    )

    add_text(doc, "4.2 语音输入和 AhaType 的关系", style="Heading 2", space_after=8)
    add_text(
        doc,
        "单纯语音转文字本身可以先独立使用。AhaType 的作用是在识别结果出来后，再做一次整理和润色，适合口语转书面语、补全标点或让内容更适合直接发送和记录。",
    )
    add_text(
        doc,
        "如果你要使用 AhaType，需要确保已经登录，并且云端服务可用。若只测试本地语音输入，不开启 AhaType 也可以正常完成基础语音转文字。",
    )

    add_text(doc, "五、模式配置与默认预设", style="Heading 1", space_after=8)
    add_text(
        doc,
        "Mac 端“模式配置”页的核心逻辑与 Windows 保持一致，这里也单独完整说明一次，避免用户因为平台不同而误解。",
    )
    add_text(doc, "5.1 模式的基本含义", style="Heading 2", space_after=8)
    add_bullets(
        doc,
        [
            "单击电源键可以切换 Mode0、Mode1、Mode2。",
            "每个模式都拥有独立的 4 个按键配置，也可以拥有独立的图片或 GIF 动画。",
            "Mode0 用于默认 Vibecoding 功能，Mode1 和 Mode2 更适合用户自定义常用快捷键、宏和显示内容。",
        ],
    )
    add_text(doc, "5.2 Mode0 默认预设", style="Heading 2", space_after=8)
    add_bullets(
        doc,
        [
            "Key1 -> F18 语音键",
            "Key2 -> YES",
            "Key3 -> NO",
            "Key4 -> Enter",
        ],
    )
    add_text(
        doc,
        "这组预设属于显示层默认值，不代表软件会在你没确认的情况下偷偷改写设备配置。如果某个键位已经有真实自定义值，界面会优先显示真实配置。",
    )
    add_tip(
        doc,
        "Key1 对应的是键盘最左边的语音键。日常测试语音输入时，优先确认自己当前切换到的是 Mode0，并按的是最左边那个语音键。",
    )

    add_text(doc, "六、如何在 Mac 上配置快捷键", style="Heading 1", space_after=8)
    add_text(
        doc,
        "下面给出一个完整的实操例子。即使 Windows 文档里已经讲过，这里也按照 Mac 的说明书逻辑再讲一遍。",
    )
    add_text(doc, "6.1 示例：把 Key1 设置为 Ctrl+C 复制", style="Heading 2", space_after=8)
    add_numbered(
        doc,
        [
            "先确保设备已经连接，并且顶部“连接”已经成功，否则后面应用到设备时会提示请先连接设备。",
            "进入“模式配置”页，选择你要修改的模式，例如 Mode1。",
            "在左侧 4 键示意图里点击 Key1。",
            "在“按键描述”里输入便于识别的名字，例如 CTRL_C_COPY。按键描述会显示在键盘屏幕上，建议使用英文、数字或下划线。",
            "在“按键类型”中选择“快捷键”。",
            "在键码下拉框中选择 Left Ctrl，点击“添加”。",
            "再选择字母 C，点击“添加”。",
            "确认列表里已经出现 Left Ctrl 和 C 两个条目后，点击“应用按键到设备”测试；如果这是完整配置的一部分，也可以最后统一执行“保存到设备”。",
        ],
    )
    add_text(
        doc,
        "快捷键适合一键触发简单组合，例如复制、粘贴、切换窗口、打开资源管理器或呼出系统功能。如果只是单次组合键，优先使用“快捷键”类型，配置和排查都会更直接。",
    )

    add_text(doc, "七、如何在 Mac 上配置宏定义", style="Heading 1", space_after=8)
    add_text(
        doc,
        "宏定义适合那些需要分步骤执行的动作。最常见的理解方式是：快捷键是“一次性组合”，宏是“按顺序执行的动作列表”。",
    )
    add_text(doc, "7.1 示例：用宏实现 Ctrl+C", style="Heading 2", space_after=8)
    add_numbered(
        doc,
        [
            "进入“模式配置”页，选中要编辑的模式，再点中目标按键，例如 Key2。",
            "在“按键描述”中输入 MACRO_COPY。",
            "在“按键类型”中选择“宏”。",
            "添加步骤“按下按键 -> Left Ctrl”。",
            "添加步骤“按下按键 -> C”。",
            "添加步骤“延时 -> 30”，约等于 90ms。",
            "添加步骤“释放按键 -> C”。",
            "添加步骤“释放按键 -> Left Ctrl”。",
            "最后添加步骤“释放全部按键”。",
            "完成后点击“应用按键到设备”或统一“保存到设备”。",
        ],
    )
    add_tip(
        doc,
        "宏的最后建议始终补一条“释放全部按键”，这样即使前面某一步释放失败，也能降低修饰键被卡住的风险。",
    )

    add_text(doc, "八、如何上传图片或 GIF 动画", style="Heading 1", space_after=8)
    add_text(
        doc,
        "“动画管理”用于给当前模式上传静态图片或 GIF，生成设备显示用的帧。Mac 端这里的工作方式也和 Windows 一致，所以同样需要单独写清楚。",
    )
    add_text(doc, "8.1 上传静态图片", style="Heading 2", space_after=8)
    add_numbered(
        doc,
        [
            "先点击顶部“连接”，确保设备配置链路已经连通。",
            "进入“模式配置”页并切换到目标模式。",
            "在右侧“动画管理”区域点击“添加图片”。",
            "选择一张或多张图片，导入后它们会出现在帧列表中。",
            "点击某一帧可以在预览区查看效果。",
            "确认无误后点击“上传到设备”。",
        ],
    )
    add_text(doc, "8.2 上传 GIF", style="Heading 2", space_after=8)
    add_numbered(
        doc,
        [
            "在“动画管理”区域点击“添加 GIF”。",
            "选择一个 GIF 文件。",
            "软件会自动把 GIF 拆成多帧并加入当前模式的帧列表。",
            "根据预览效果调整 FPS。FPS 越高，动画播放越快。",
            "确认效果后点击“上传到设备”，最后如有需要再执行“保存到设备”。",
        ],
    )
    add_tip(
        doc,
        "如果动画管理里点上传或保存时提示请先连接设备，仍然说明是顶部配置通信链路没有连好，而不是图片文件本身有问题。",
    )

    add_text(doc, "九、Mac 端常见问题", style="Heading 1", space_after=8)
    faq_items = [
        "问题 1：蓝牙已经连上了，为什么还是不能把按键应用到设备？\n答：蓝牙连接和上位机配置连接不是同一件事。请先确认顶部“连接”已经成功，再进行应用或保存。",
        "问题 2：点了“启动语音输入”但没有反应。\n答：先看顶部状态是否最终变成“语音已就绪”。如果没有，优先检查麦克风、输入监控、辅助功能三项权限是否都已打开。",
        "问题 3：Mac 上第一次打开为什么会连续弹出很多权限框？\n答：这是正常现象。首次使用需要把文稿访问、蓝牙、麦克风、输入监控、辅助功能以及 Hook 安装这几件事补齐，后面就会稳定很多。",
        "问题 4：Hook 一定要装吗？\n答：如果你要完整体验 Claude / Cursor 联动，建议安装；如果你只是先测试基础语音输入，Hook 不是唯一阻断项，但也建议第一次就补齐。",
        "问题 5：Mode0 看起来像已经有预设，是不是已经写进设备了？\n答：不一定。Mode0 的 F18、YES、NO、Enter 默认显示属于界面层预设，用来告诉用户默认含义，不等于软件在后台静默改写了真实配置。",
    ]
    for item in faq_items:
        add_text(doc, item)

    add_text(doc, "十、给用户的最快上手路径", style="Heading 1", space_after=8)
    add_numbered(
        doc,
        [
            "先把应用拖到 Applications，再从“应用程序”里启动。",
            "允许蓝牙和麦克风权限，并按提示完成 Hook 安装。",
            "点击“启动语音输入”，继续完成输入监控和辅助功能授权。",
            "等状态变成“语音已就绪”后，把光标放到输入框里，按住键盘最左边的语音键说话。",
            "如果要改快捷键、宏或动画，再去“模式配置”页，并记得先点击顶部“连接”。",
        ],
    )
    add_tip(
        doc,
        "这份说明书是 Mac 独立版，所以这里把 Windows 已解释过的概念重新讲了一遍。你给用户发文档时，优先发这一份即可，不需要再额外搭配 Windows 版让用户自己拼流程。",
    )


def main() -> None:
    doc = Document()
    configure_document(doc)
    build_manual(doc)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
